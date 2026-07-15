import os
import csv
import shutil
import importlib
from pathlib import Path
from typing import List, Dict, Any, Optional

from assistant.config import (
    RAG_ENABLED,
    RAG_DOCUMENTS_DIR,
    RAG_VECTOR_DB_DIR,
    RAG_CHUNK_SIZE,
    RAG_CHUNK_OVERLAP,
    RAG_NUM_RETRIEVED,
)
from assistant.utils import log, log_system_metrics
from backend.services.hybrid_retrieval import reciprocal_rank_fusion
import time

class RAGService:
    def __init__(self):
        self.vectorstore = None
        self._reranker = None
        self._lexical_word_vectorizer = None
        self._lexical_word_matrix = None
        self._lexical_char_vectorizer = None
        self._lexical_char_matrix = None
        self._lexical_documents = []
        if RAG_ENABLED:
            self._load_vectorstore()
            existing_count = self._vectorstore_count()
            if not self.vectorstore or existing_count == 0:
                try:
                    docs_dir = Path(RAG_DOCUMENTS_DIR)
                    if not docs_dir.exists():
                        log(f"RAG documents directory not found: {RAG_DOCUMENTS_DIR}", title="RAG", style="bold yellow")
                    else:
                        supported_count = len(self._iter_document_paths())
                        if supported_count == 0:
                            log(f"No supported documents found in {RAG_DOCUMENTS_DIR}; skipping RAG build.", title="RAG", style="bold yellow")
                        else:
                            log(f"Building RAG vector DB from {RAG_DOCUMENTS_DIR} ({supported_count} document(s)).", title="RAG", style="bold blue")
                            self._load_documents_into_vectorstore()
                except Exception as exc:
                    log(f"Automatic RAG build failed: {exc}", title="RAG", style="bold yellow")

    def _embedding_device(self) -> str:
        try:
            torch = importlib.import_module("torch")
            return "cuda" if torch.cuda.is_available() else "cpu"
        except Exception:
            return "cpu"

    def _supported_document_suffixes(self) -> set[str]:
        return {".pdf", ".txt", ".md", ".csv", ".docx"}

    def _iter_document_paths(self) -> List[Path]:
        docs_dir = Path(RAG_DOCUMENTS_DIR)
        if not docs_dir.exists():
            return []
        return [p for p in docs_dir.rglob("*") if p.is_file() and p.suffix.lower() in self._supported_document_suffixes()]

    def _read_document_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".csv":
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                reader = csv.reader(handle)
                rows = [", ".join(row) for row in reader]
            return "\n".join(rows)
        if suffix == ".docx":
            try:
                from docx import Document
                doc = Document(str(path))
                return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            except Exception as exc:
                log(f"Skipping DOCX {path.name}: {exc}", title="RAG", style="bold yellow")
                return ""
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                pages = []
                for page in reader.pages:
                    try:
                        pages.append(page.extract_text() or "")
                    except Exception:
                        pages.append("")
                return "\n".join(page for page in pages if page.strip())
            except Exception as exc:
                log(f"Skipping PDF {path.name}: {exc}", title="RAG", style="bold yellow")
                return ""
        return ""

    def _chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        clean_text = text.strip()
        if not clean_text:
            return []
        if chunk_size <= 0:
            return [clean_text]
        chunks = []
        step = max(1, chunk_size - max(0, chunk_overlap))
        for start in range(0, len(clean_text), step):
            chunk = clean_text[start:start + chunk_size].strip()
            if chunk:
                chunks.append(chunk)
        return chunks

    def _format_documents(self, docs: List[Any]) -> str:
        """Format untrusted evidence with stable, user-visible citations."""
        parts = []
        for doc in docs:
            metadata = getattr(doc, "metadata", {}) or {}
            source = os.path.basename(str(metadata.get("source", "unknown")))
            page = metadata.get("page")
            location = f"page {page}" if page is not None else f"chunk {metadata.get('chunk', '?')}"
            citation = f"[{source}, {location}]"
            parts.append(f"{citation}\n{getattr(doc, 'page_content', str(doc))}")
        return "\n\n".join(parts)

    def _rerank_documents(self, query: str, docs: List[Any], top_n: int = 5) -> List[Any]:
        if not docs:
            return []
        pairs = [(query, getattr(doc, "page_content", str(doc))) for doc in docs]
        scores = []
        try:
            if self._reranker is None:
                from sentence_transformers import CrossEncoder
                self._reranker = CrossEncoder("BAAI/bge-reranker-base", device=self._embedding_device())
            raw_scores = self._reranker.predict(pairs)
            scores = [float(score) for score in raw_scores]
        except Exception:
            query_terms = {token for token in query.lower().split() if len(token) > 2}
            for doc in docs:
                text = getattr(doc, "page_content", str(doc)).lower()
                doc_terms = {token for token in text.split() if len(token) > 2}
                overlap = len(query_terms & doc_terms)
                scores.append(float(overlap))
        ranked = sorted(zip(docs, scores), key=lambda item: item[1], reverse=True)
        top_docs = []
        for doc, score in ranked[:top_n]:
            meta = dict(getattr(doc, "metadata", {}) or {})
            meta["score"] = score
            try:
                doc.metadata = meta
            except Exception:
                pass
            top_docs.append(doc)
        return top_docs

    def _build_lexical_index(self, texts: List[str], metadatas: List[dict]) -> None:
        """Build the word + character TF-IDF side of hybrid retrieval."""
        if not texts:
            self._lexical_documents = []
            return
        try:
            from langchain_core.documents import Document
            from sklearn.feature_extraction.text import TfidfVectorizer

            self._lexical_word_vectorizer = TfidfVectorizer(
                lowercase=True, ngram_range=(1, 2), sublinear_tf=True
            )
            self._lexical_char_vectorizer = TfidfVectorizer(
                analyzer="char_wb", lowercase=True, ngram_range=(3, 5), min_df=2, sublinear_tf=True
            )
            self._lexical_word_matrix = self._lexical_word_vectorizer.fit_transform(texts)
            self._lexical_char_matrix = self._lexical_char_vectorizer.fit_transform(texts)
            self._lexical_documents = [
                Document(page_content=text, metadata=dict(metadata))
                for text, metadata in zip(texts, metadatas)
            ]
            log(f"Lexical index built with {len(texts)} chunk(s).", title="RAG", style="bold green")
        except Exception as exc:
            self._lexical_documents = []
            log(f"Lexical index unavailable: {exc}", title="RAG", style="bold yellow")

    def _refresh_lexical_index_from_vectorstore(self) -> None:
        if not self.vectorstore:
            return
        try:
            payload = getattr(self.vectorstore, "_collection").get(include=["documents", "metadatas"])
            texts = list(payload.get("documents") or [])
            metadatas = list(payload.get("metadatas") or [{} for _ in texts])
            self._build_lexical_index(texts, metadatas)
        except Exception as exc:
            log(f"Could not restore lexical index: {exc}", title="RAG", style="bold yellow")

    def _lexical_search(self, query: str, k: int, source_filter: Optional[str] = None) -> List[Any]:
        if not self._lexical_documents or self._lexical_word_vectorizer is None:
            return []
        try:
            import numpy as np

            word_query = self._lexical_word_vectorizer.transform([query])
            char_query = self._lexical_char_vectorizer.transform([query])
            word_scores = (word_query @ self._lexical_word_matrix.T).toarray()[0]
            char_scores = (char_query @ self._lexical_char_matrix.T).toarray()[0]
            scores = 0.65 * word_scores + 0.35 * char_scores
            order = np.argsort(-scores, kind="stable")
            filter_name = os.path.basename(source_filter).lower() if source_filter else None
            results = []
            for index in order:
                document = self._lexical_documents[int(index)]
                source = os.path.basename(str(document.metadata.get("source", ""))).lower()
                if filter_name and filter_name not in source:
                    continue
                metadata = dict(document.metadata)
                metadata["lexical_score"] = float(scores[int(index)])
                document.metadata = metadata
                results.append(document)
                if len(results) >= k:
                    break
            return results
        except Exception as exc:
            log(f"Lexical retrieval failed: {exc}", title="RAG", style="bold yellow")
            return []
    def search_knowledge_base(self, query: str, top_k: int = 10, rerank_k: int = 5, source_filter: Optional[str] = None) -> str:
        if not self.vectorstore:
            return ""
        search_query = query.strip() or query
        candidate_k = max(top_k * 4, 20)

        dense_started = time.time()
        try:
            dense_docs = self.vectorstore.similarity_search(search_query, k=candidate_k)
        except Exception:
            try:
                retriever = self.vectorstore.as_retriever(search_kwargs={"k": candidate_k})
                dense_docs = retriever.get_relevant_documents(search_query)
            except Exception as exc:
                log(f"Dense retrieval failed: {exc}", title="RAG", style="bold yellow")
                dense_docs = []
        dense_ms = (time.time() - dense_started) * 1000

        if source_filter:
            filter_name = os.path.basename(source_filter).lower()
            dense_docs = [
                doc for doc in dense_docs
                if filter_name in os.path.basename(str(doc.metadata.get("source", ""))).lower()
            ]

        lexical_started = time.time()
        lexical_docs = self._lexical_search(search_query, candidate_k, source_filter=source_filter)
        lexical_ms = (time.time() - lexical_started) * 1000

        fused_docs = reciprocal_rank_fusion(
            [dense_docs, lexical_docs], limit=candidate_k, rank_constant=10, weights=[0.35, 1.0]
        )
        if not fused_docs:
            log("No documents retrieved from dense or lexical index.", title="RAG", style="bold yellow")
            return ""

        use_reranker = os.getenv("RAG_ENABLE_RERANKER", "false").strip().lower() in {"1", "true", "yes"}
        rerank_ms = 0.0
        if use_reranker:
            rerank_started = time.time()
            final_docs = self._rerank_documents(search_query, fused_docs, top_n=rerank_k)
            rerank_ms = (time.time() - rerank_started) * 1000
        else:
            final_docs = fused_docs[:rerank_k]
        log(
            f"Hybrid retrieval: dense={len(dense_docs)} ({dense_ms:.1f} ms), "
            f"lexical={len(lexical_docs)} ({lexical_ms:.1f} ms), "
            f"fused={len(fused_docs)}, reranker={use_reranker} ({rerank_ms:.1f} ms)",
            title="Metrics: RAG",
            style="bold cyan",
        )
        return self._format_documents(final_docs)

    def _load_documents_into_vectorstore(self) -> None:
        start_time = time.time()
        try:
            HuggingFaceEmbeddings = importlib.import_module("langchain_huggingface").HuggingFaceEmbeddings
            Chroma = importlib.import_module("langchain_chroma").Chroma
        except Exception as exc:
            log(f"RAG dependencies unavailable: {exc}", title="RAG", style="bold yellow")
            return
        docs_dir = Path(RAG_DOCUMENTS_DIR)
        if not docs_dir.exists():
            log(f"RAG documents directory missing: {RAG_DOCUMENTS_DIR}", title="RAG", style="bold yellow")
            return
        documents = []
        for path in self._iter_document_paths():
            text = self._read_document_text(path)
            if not text.strip():
                continue
            for index, chunk in enumerate(self._chunk_text(text, RAG_CHUNK_SIZE, RAG_CHUNK_OVERLAP), start=1):
                documents.append({"text": chunk, "source": str(path), "chunk": str(index), "type": path.suffix.lower().lstrip(".")})
        if not documents:
            log("No supported documents found for RAG build.", title="RAG", style="bold yellow")
            self.vectorstore = None
            return

        chunking_time = time.time() - start_time
        log(f"Document Parsing & Chunking Time: {chunking_time:.3f} s", title="Metrics: RAG", style="bold cyan")

        ingestion_start = time.time()
        shutil.rmtree(RAG_VECTOR_DB_DIR, ignore_errors=True)
        os.makedirs(RAG_VECTOR_DB_DIR, exist_ok=True)
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"device": self._embedding_device()})
        texts = [item["text"] for item in documents]
        metadatas = [{"source": item["source"], "chunk": item["chunk"], "type": item["type"]} for item in documents]
        self.vectorstore = Chroma.from_texts(texts=texts, embedding=embeddings, metadatas=metadatas, persist_directory=RAG_VECTOR_DB_DIR)
        self._build_lexical_index(texts, metadatas)

        ingestion_time = time.time() - ingestion_start
        total_time = time.time() - start_time
        log(f"Built RAG vectorstore with {len(texts)} chunk(s). Ingestion Time: {ingestion_time:.3f} s | Total: {total_time:.3f} s", title="RAG", style="bold green")

        log_system_metrics("Metrics: Post-Ingestion System Resources")

    def _load_vectorstore(self) -> None:
        try:
            HuggingFaceEmbeddings = importlib.import_module("langchain_huggingface").HuggingFaceEmbeddings
            Chroma = importlib.import_module("langchain_chroma").Chroma
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"device": self._embedding_device()})
            self.vectorstore = Chroma(persist_directory=RAG_VECTOR_DB_DIR, embedding_function=embeddings)
            self._refresh_lexical_index_from_vectorstore()
            count = None
            try:
                count = getattr(self.vectorstore, "_collection").count()
            except Exception:
                pass
            log(f"RAG vectorstore loaded (dir={RAG_VECTOR_DB_DIR}) count={count}", title="RAG", style="bold green")
        except Exception as exc:
            self.vectorstore = None
            log(f"Could not load RAG vectorstore: {exc}", title="RAG", style="bold yellow")

    def _vectorstore_count(self) -> int:
        if not self.vectorstore:
            return 0
        try:
            count = getattr(self.vectorstore, "_collection").count()
            return int(count or 0)
        except Exception:
            return 0

    def rebuild_rag(self) -> None:
        try:
            log("Rebuilding RAG vector DB with supported document formats...", title="RAG", style="bold blue")
            self._load_documents_into_vectorstore()
            log("RAG vectorstore updated successfully.", title="RAG", style="bold green")
        except Exception as exc:
            log(f"Dynamic RAG rebuild failed: {exc}", title="RAG", style="bold red")

    def retrieve_and_process(self, query: str, source_filter: Optional[str] = None) -> str:
        return self.search_knowledge_base(
            query=query,
            top_k=RAG_NUM_RETRIEVED,
            rerank_k=RAG_NUM_RETRIEVED,
            source_filter=source_filter,
        )

    def should_use_knowledge_base(self, prompt: str, document: Optional[dict] = None) -> bool:
        if not self.vectorstore or self._vectorstore_count() == 0:
            return False
        if len((prompt or "").strip()) < 5:
            return False
        if not document:
            return False
        return True
